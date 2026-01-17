import os
import io
import re
import logging
import PyPDF2
from typing import Optional
from dataclasses import dataclass

logger = logging.getLogger("document_service")

@dataclass
class ProcessedDocument:
    content: str
    filename: str
    doc_type: str
    word_count: int
    metadata: dict

class DocumentService:
    """
    Enterprise Document Processing Service
    Handles PDF, DOCX, and Text extraction with robust error handling.
    """
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Normalize whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip noise (page numbers, very short lines)
            if len(line) < 3 or line.isdigit():
                continue
            cleaned_lines.append(line)
            
        return '\n'.join(cleaned_lines)

    def _upload_to_s3(self, content: bytes, filename: str) -> str:
        """Upload file to AWS S3 and return the URL"""
        try:
            import boto3
            # Reload env check
            if not os.environ.get("AWS_ACCESS_KEY_ID"):
                from dotenv import load_dotenv
                load_dotenv()

            # Prefer os.environ for Boto3 as it honors it natively usually, 
            # but we will pass explicitly to be safe.
            aws_key = os.environ.get("AWS_ACCESS_KEY_ID")
            aws_secret = os.environ.get("AWS_SECRET_ACCESS_KEY")
            aws_region = os.environ.get("AWS_REGION", "eu-central-1")
            bucket = os.environ.get("S3_BUCKET_NAME")

            logger.info(f"S3 Connection Attempt: Key={aws_key[:4]}*** Bucket={bucket} Region={aws_region}")

            if not aws_key or not aws_secret:
                raise ValueError("Missing AWS Credentials in Environment")

            s3_client = boto3.client(
                's3',
                aws_access_key_id=aws_key,
                aws_secret_access_key=aws_secret,
                region_name=aws_region
            )
            
            # Use Content Type based on extension
            content_type = 'application/octet-stream'
            if filename.endswith('.pdf'): content_type = 'application/pdf'
            elif filename.endswith('.docx'): content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            s3_client.put_object(
                Bucket=bucket,
                Key=f"documents/{filename}",
                Body=content,
                ContentType=content_type
            )
            
            return f"s3://{bucket}/documents/{filename}"
            
        except Exception as e:
            logger.error(f"S3 Upload failed: {e}")
            return f"local://{filename}" # Fallback
            
    async def process_file(self, content: bytes, filename: str, skip_upload: bool = False) -> Optional[ProcessedDocument]:
        """Process file content based on extension"""
        ext = filename.lower().split('.')[-1]
        extracted_text = ""
        
        try:
            # 1. Upload to Cloud (S3) for Persistence (Optional)
            if not skip_upload:
                storage_url = self._upload_to_s3(content, filename)
                logger.info(f"File stored at: {storage_url}")
            else:
                # Assume it's already in S3 (Sync Mode)
                bucket = os.environ.get("S3_BUCKET_NAME", "unknown-bucket")
                storage_url = f"s3://{bucket}/documents/{filename}"
                logger.info(f"Skipping S3 Upload (Sync Mode). Using: {storage_url}")
            
            if ext == 'pdf':
                extracted_text = self._extract_pdf(content)
            elif ext == 'docx':
                extracted_text = self._extract_docx(content)
            elif ext in ['txt', 'md']:
                extracted_text = content.decode('utf-8', errors='ignore')
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return None
                
            clean_content = self.clean_text(extracted_text)
            
            if not clean_content:
                logger.warning(f"No content extracted from {filename}")
                return None
                
            return ProcessedDocument(
                content=clean_content,
                filename=filename,
                doc_type=ext,
                word_count=len(clean_content.split()),
                metadata={
                    "source": storage_url,  # NOW USING S3 URL
                    "processed_at": str(os.times())
                }
            )

        except Exception as e:
            logger.error(f"Failed to process {filename}: {str(e)}")
            return None

    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF bytes"""
        text_content = ""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content += f"\n\n--- Page {i+1} ---\n\n{page_text}"
            return text_content
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            logger.error("python-docx not installed")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise

document_service = DocumentService()
